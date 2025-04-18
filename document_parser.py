import os
import json
import csv
from typing import Dict, List, Any, Union, Optional
from pdf2image import convert_from_path
import cv2
import numpy as np
import pytesseract
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
from openpyxl import load_workbook
from docx.enum.shape import WD_INLINE_SHAPE
from PIL import Image
import logging
from concurrent.futures import ProcessPoolExecutor, ThreadPoolExecutor
from dataclasses import dataclass
from pathlib import Path
import multiprocessing
from functools import partial
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@dataclass
class ParsedDocument:
    """Data class to store parsed document information"""
    file_path: str
    content: str
    metadata: Dict[str, Any]
    success: bool
    error_message: Optional[str] = None

class DocumentParser:
    """Main parser class that handles multiple document types and parallel processing"""
    
    def __init__(self, num_workers: int = None, use_process_pool: bool = True):
        """
        Initialize the parser with configuration options
        
        Args:
            num_workers: Number of parallel workers (defaults to CPU count)
            use_process_pool: If True, uses ProcessPoolExecutor, else ThreadPoolExecutor
        """
        self.num_workers = num_workers or multiprocessing.cpu_count()
        self.use_process_pool = use_process_pool
        self.supported_formats = {
            '.pdf': self._parse_pdf,
            '.html': self._parse_html,
            '.docx': self._parse_docx,
            '.xlsx': self._parse_xlsx,
            '.png': self._parse_image,
            '.jpg': self._parse_image,
            '.jpeg': self._parse_image
        }

    @staticmethod
    def _preprocess_image(image: np.ndarray) -> np.ndarray:
        """Preprocess image for better text extraction"""
        try:
            if len(image.shape) == 3:
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            else:
                gray = image
            blurred = cv2.GaussianBlur(gray, (5, 5), 0)
            _, thresh = cv2.threshold(blurred, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            return thresh
        except Exception as e:
            logger.error(f"Error preprocessing image: {str(e)}")
            return image

    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """Parse PDF documents"""
        try:
            images = convert_from_path(file_path,poppler_path=r"C:\Program Files\poppler-24.08.0\Library\bin")
            extracted_text = []
            
            for i, img in enumerate(images):
                processed_image = self._preprocess_image(np.array(img))
                page_text = pytesseract.image_to_string(processed_image)
                extracted_text.append(f"Page {i + 1}:\n{page_text}")
            
            return ParsedDocument(
                file_path=file_path,
                content='\n\n'.join(extracted_text),
                metadata={'num_pages': len(images)},
                success=True
            )
        except Exception as e:
            return ParsedDocument(
                file_path=file_path,
                content='',
                metadata={},
                success=False,
                error_message=str(e)
            )

    def _parse_html(self, file_path: str) -> ParsedDocument:
        """Parse HTML documents"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            soup = BeautifulSoup(html_content, 'html.parser')
            return ParsedDocument(
                file_path=file_path,
                content=soup.get_text(),
                metadata={'title': soup.title.string if soup.title else None},
                success=True
            )
        except Exception as e:
            return ParsedDocument(
                file_path=file_path,
                content='',
                metadata={},
                success=False,
                error_message=str(e)
            )

    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse DOCX documents"""
        try:
            doc = Document(file_path)
            full_text = []
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    full_text.append(' | '.join(cell.text for cell in row.cells))

            return ParsedDocument(
                file_path=file_path,
                content='\n'.join(full_text),
                metadata=metadata,
                success=True
            )
        except Exception as e:
            return ParsedDocument(
                file_path=file_path,
                content='',
                metadata={},
                success=False,
                error_message=str(e)
            )

    def _parse_xlsx(self, file_path: str) -> ParsedDocument:
        """Parse XLSX documents"""
        try:
            wb = load_workbook(file_path, data_only=True)
            data = []
            metadata = {'sheets': wb.sheetnames}
            
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                sheet_data = []
                for row in ws.iter_rows(values_only=True):
                    sheet_data.append([str(cell) if cell is not None else '' for cell in row])
                data.append(f"Sheet: {sheet}\n" + '\n'.join(['\t'.join(row) for row in sheet_data]))
            
            return ParsedDocument(
                file_path=file_path,
                content='\n\n'.join(data),
                metadata=metadata,
                success=True
            )
        except Exception as e:
            return ParsedDocument(
                file_path=file_path,
                content='',
                metadata={},
                success=False,
                error_message=str(e)
            )

    def _parse_image(self, file_path: str) -> ParsedDocument:
        """Parse image files"""
        try:
            image = cv2.imread(file_path)
            processed_image = self._preprocess_image(image)
            extracted_text = pytesseract.image_to_string(processed_image)
            
            metadata = {
                'dimensions': image.shape[:2],
                'channels': image.shape[2] if len(image.shape) > 2 else 1
            }
            
            return ParsedDocument(
                file_path=file_path,
                content=extracted_text,
                metadata=metadata,
                success=True
            )
        except Exception as e:
            return ParsedDocument(
                file_path=file_path,
                content='',
                metadata={},
                success=False,
                error_message=str(e)
            )

    def _parse_single_file(self, file_path: str) -> ParsedDocument:
        """Parse a single file based on its extension"""
        try:
            file_extension = Path(file_path).suffix.lower()
            if file_extension not in self.supported_formats:
                return ParsedDocument(
                    file_path=file_path,
                    content='',
                    metadata={},
                    success=False,
                    error_message=f"Unsupported file format: {file_extension}"
                )
            
            parser_func = self.supported_formats[file_extension]
            return parser_func(file_path)
            
        except Exception as e:
            return ParsedDocument(
                file_path=file_path,
                content='',
                metadata={},
                success=False,
                error_message=str(e)
            )

    def parse_files(self, file_paths: List[str]) -> List[ParsedDocument]:
        """
        Parse multiple files in parallel
        
        Args:
            file_paths: List of paths to files to parse
            
        Returns:
            List of ParsedDocument objects containing the parsing results
        """
        executor_class = ProcessPoolExecutor if self.use_process_pool else ThreadPoolExecutor
        
        with executor_class(max_workers=self.num_workers) as executor:
            results = list(executor.map(self._parse_single_file, file_paths))
        
        return results

    @property
    def supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return list(self.supported_formats.keys())

